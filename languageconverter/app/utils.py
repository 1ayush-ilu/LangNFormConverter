import os
import uuid
from langdetect import detect
from deep_translator import GoogleTranslator
from PyPDF2 import PdfReader
from docx import Document

# --- Text extraction ---
def extract_text_from_file(filepath: str) -> str:
    _, ext = os.path.splitext(filepath.lower())

    if ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])

    if ext == ".pdf":
        text = []
        with open(filepath, "rb") as f:
            reader = PdfReader(f)
            for page in reader.pages:
                try:
                    text.append(page.extract_text() or "")
                except Exception:
                    text.append("")
        return "\n".join(text).strip()

    raise ValueError("Unsupported file type. Please upload TXT, DOCX, or PDF.")

# --- Language detection ---
def detect_language(text: str) -> str:
    text_sample = (text or "")[:2000]
    if not text_sample.strip():
        return "auto"
    try:
        return detect(text_sample)
    except Exception:
        return "auto"

# --- Translation (with graceful fallback) ---
def translate_text(text: str, target_lang: str) -> str:
    target_lang = (target_lang or "en").strip().lower()
    if not text.strip():
        return ""

    try:
        translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
        if isinstance(translated, list):
            translated = "\n".join(translated)
        return translated
    except Exception:
        # Fallback: no translation, return original so pipeline still works
        return text

# --- File writers ---
def save_as_txt(text: str, out_dir: str) -> str:
    filename = f"converted_{uuid.uuid4().hex[:8]}.txt"
    path = os.path.join(out_dir, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path

def save_as_docx(text: str, out_dir: str) -> str:
    filename = f"converted_{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(out_dir, filename)
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)
    return path

def save_as_pdf(text: str, out_dir: str) -> str:
    # Simple text->PDF using reportlab
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm

    filename = f"converted_{uuid.uuid4().hex[:8]}.pdf"
    path = os.path.join(out_dir, filename)

    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4

    # Basic text wrapping
    margin = 20 * mm
    max_width = width - 2 * margin
    y = height - margin
    line_height = 12

    def wrap_text(s, max_chars=95):
        # rough wrap by characters for simplicity
        lines = []
        for para in s.splitlines():
            if len(para) <= max_chars:
                lines.append(para)
            else:
                start = 0
                while start < len(para):
                    lines.append(para[start:start+max_chars])
                    start += max_chars
        return lines

    c.setFont("Helvetica", 11)
    for line in wrap_text(text):
        if y < margin + line_height:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - margin
        c.drawString(margin, y, line)
        y -= line_height

    c.save()
    return path
